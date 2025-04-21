import numpy as np
import matplotlib.pyplot as plt
from scipy.signal import freqz
from docx import Document
from docx.shared import Inches
import io

# Süzgeç parametreleri
L = 5
h = np.ones(L) / L  # Impuls cevabı

# Frekans cevabını hesapla (freqz)
w, H = freqz(h, worN=1000, whole=False)
w_full = np.linspace(-np.pi, np.pi, 1000)
H_analytic = (1/L) * np.sin(w_full * L/2) / (np.sin(w_full/2) + 1e-10) * np.exp(-1j * w_full * (L-1)/2)

# Grafikleri oluştur
plt.figure(figsize=(12, 6))

# Genlik Cevabı
plt.subplot(2, 1, 1)
plt.plot(w_full, np.abs(H_analytic))
plt.title("Genlik Cevabı")
plt.xlabel(r"$\omega$")
plt.ylabel(r"$|H(\omega)|$")
plt.grid(True)

# Faz Cevabı
plt.subplot(2, 1, 2)
plt.plot(w_full, np.angle(H_analytic))
plt.title("Faz Cevabı")
plt.xlabel(r"$\omega$")
plt.ylabel(r"$\angle H(\omega)$")
plt.grid(True)

plt.tight_layout()

# Grafikleri belleğe kaydet
img_buffer = io.BytesIO()
plt.savefig(img_buffer, format='png', dpi=300)
img_buffer.seek(0)

# Word belgesi oluştur
doc = Document()
doc.add_heading('Süzgeç Frekans Cevabı Analizi', level=1)

# Bilgi ekle
doc.add_paragraph(
    f"Impuls cevabı: h[n] = 1/{L} için n = 0, 1, ..., {L-1}\n"
    "Frekans cevabı DTFT ile hesaplanmıştır."
)

# Grafikleri ekle
doc.add_heading('Genlik ve Faz Cevabı', level=2)
doc.add_picture(img_buffer, width=Inches(6.0))

# Analiz sonuçları
doc.add_heading('Analiz Sonuçları', level=2)
doc.add_paragraph(
    "1. Genlik cevabı alçak geçiren (low-pass) karakteristik gösterir.\n"
    f"2. Sıfırlar: ω = ±2πk/{L} (k = 1, 2, ..., {L-1})\n"
    f"3. Faz cevabı lineerdir: ∠H(ω) = -ω*{(L-1)/2}\n"
    "4. Bu bir FIR (sonlu impuls cevabı) süzgeçtir."
)

# Dosyayı kaydet
doc.save('süzgeç_frekans_cevabı.docx')
print("Word dosyası oluşturuldu: süzgeç_frekans_cevabı.docx")